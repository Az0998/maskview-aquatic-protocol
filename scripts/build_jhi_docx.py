"""Build IWA Journal of Hydroinformatics Word manuscript (Harvard refs)."""
from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
PAPER = ROOT / "paper"
TABLES = ROOT / "results" / "tables"
FIGS = ROOT / "results" / "figures"
LAKE = ROOT / "data" / "frozen" / "lake"
OCEAN = ROOT / "data" / "frozen" / "ocean"


def font(run, size=11, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def h(doc, text, level=1):
    p = doc.add_paragraph()
    font(p.add_run(text), size=12 if level == 1 else 11, bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def body(doc, text):
    p = doc.add_paragraph()
    font(p.add_run(text), size=11)
    p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE


def center_author(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Senjie Zhang")
    font(r, size=11)
    s = p.add_run("1,*")
    font(s, size=11)
    s.font.superscript = True
    p.paragraph_format.space_after = Pt(6)


def center(doc, text, size=11, italic=False, after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(text), size=size, italic=italic)
    p.paragraph_format.space_after = Pt(after)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(p.add_run(text), size=10, italic=True)
    p.paragraph_format.space_after = Pt(10)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    for i, ht in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        font(cell.paragraphs[0].add_run(str(ht)), size=9, bold=True)
    for r_i, row in enumerate(rows):
        for c_i, val in enumerate(row):
            cell = t.rows[r_i + 1].cells[c_i]
            cell.text = ""
            font(cell.paragraphs[0].add_run(str(val)), size=9)
    return t


def maybe_fig(doc, path: Path, width=6.1) -> bool:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        return True
    return False


def f3(x):
    return f"{float(x):.3f}"


def mae_at(df, protocol, setting, model):
    sub = df[(df.protocol == protocol) & (df.setting == setting) & (df.model == model)]
    return f3(sub.MAE.iloc[0]) if not sub.empty else "—"


def keep_row(df, **kwargs):
    sub = df.copy()
    for k, v in kwargs.items():
        sub = sub[sub[k] == v]
    r = sub.iloc[0]
    return [
        f3(r.keep_frac),
        f3(r.st_rmse),
        f3(r.point_delta_clim_minus_st),
        f"[{f3(r.delta_clim_minus_st_p05)}, {f3(r.delta_clim_minus_st_p95)}]",
    ]


def main():
    lake = pd.read_csv(TABLES / "lake_pattern_winners.csv")
    ocean = pd.read_csv(TABLES / "ocean_pattern_winners.csv")
    rr = pd.read_csv(TABLES / "rank_reversal_shared_patterns.csv")
    eco = pd.read_csv(LAKE / "ecoinf_dual_protocol.csv")
    keep = pd.read_csv(OCEAN / "keep_ratio_scan.csv")
    boot = pd.read_csv(OCEAN / "st_clim_bootstrap.csv")
    wins = json.loads((LAKE / "paper_tables_summary.json").read_text(encoding="utf-8"))
    hl_path = PAPER / "highlights.txt"
    highlights = [
        ln.lstrip("• ").strip()
        for ln in hl_path.read_text(encoding="utf-8").splitlines()
        if ln.strip()
    ]

    doc = Document()
    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font(
        title.add_run(
            "Operational missingness reverses method rankings in aquatic monitoring "
            "networks: a Mask-View protocol on a lake station grid and a shelf oxygen cube"
        ),
        size=14,
        bold=True,
    )
    center_author(doc)
    center(doc, "1 Lanzhou University, Lanzhou 730000, Gansu, China", size=10, italic=True)
    center(doc, "*Corresponding author: 3079099853@qq.com", size=10, italic=True)
    center(
        doc,
        "Journal of Hydroinformatics (IWA)  ·  Research paper  ·  Harvard references",
        size=9,
        italic=True,
        after=10,
    )

    h(doc, "Highlights")
    for bullet in highlights:
        p = doc.add_paragraph(style="List Bullet")
        font(p.add_run(bullet), size=11)

    h(doc, "Abstract")
    body(
        doc,
        "Automatic water-quality stations and biogeochemical profiles rarely fail as scattered "
        "missing-completely-at-random (MCAR) holes. They fail as downtime slabs, dropped parameters, "
        "and station or Argo-column blackouts. Benchmarks that hide only random cells therefore overstate "
        "the value of local temporal interpolators and of models trained on dense history. We introduce a "
        "shared Mask-View pattern bank (point, block_time, sensor, station, mixed, and an Argo-column analog) "
        "and score who wins, not a universal error. On the Dianchi 22-station, nine-variable, 4-hourly panel "
        "(2022–2024), linear interpolation attains the lowest standardized mean absolute error (MAE) on point "
        "(0.161), block_time (0.254) and mixed (0.264) gaps, but collapses to about 0.68–0.70 under sensor and "
        "station outages, where BRITS and a MixHop spatiotemporal imputer recover. Linear still wins a 2025 "
        "Ecological Informatics random/week-gap bank and a year-shift replay of real outages; ranking reverses "
        "only when the evaluation window has no temporal neighbours. On an East China Sea shelf dissolved-oxygen "
        "cube, a spatiotemporal Transformer is best at one-month lead under dense history (root-mean-square error "
        "3.88 µmol kg−1 versus climatology 5.30). After fair masked persistence is scored, the Transformer remains "
        "first at lead-1; the default station/section margin over climatology shrinks to about 0.06 µmol kg−1 "
        "mainly because those masks keep 8–9% of voxels. A paired month-block bootstrap still excludes zero. "
        "Lead-2 skill reverts to climatology except for block_time. MCAR rankings are not operational rankings. "
        "Hydroinformatics papers should report an operational pattern bank before claiming that a learned model "
        "replaces linear interpolation or climatology.",
    )

    h(doc, "Keywords")
    p = doc.add_paragraph()
    font(
        p.add_run(
            "missing data; water quality; dissolved oxygen; graph neural network; hydroinformatics; "
            "Mask-View; automatic monitoring; East China Sea; Dianchi Lake; imputation"
        ),
        size=11,
    )

    h(doc, "1. Introduction")
    body(
        doc,
        "High-frequency aquatic monitoring is now limited less by sampling design than by structured "
        "missingness (Little & Rubin 2019; van Buuren 2018). Instruments go offline for hours to days; "
        "a single parameter channel dies; a whole station or profile column disappears. National automatic "
        "networks in China deliver multi-parameter series at sub-daily resolution (CNEMC 2026), creating "
        "both an early-warning opportunity and a missing-data problem. On the open Dianchi automatic-station "
        "release the natural missing rate on a regular 4 h grid is 19.8%. Of those missing cells, 96.9% occur "
        "while the entire station is silent; isolated one-step holes are 0.6%. Station-outage durations are "
        "heavy-tailed (cell-weighted median 3.5 days). Ocean oxygen products face the complementary problem: "
        "biogeochemical Argo profiles sample the interior as sparse columns, not as a dense cube "
        "(Sharp et al. 2023; Breitburg et al. 2018).",
    )
    body(
        doc,
        "Deep imputers such as BRITS (Cao et al. 2018) and SAITS (Du et al. 2023), recurrent models for "
        "missing values (Che et al. 2018), generative imputers (Luo et al. 2018), geo-sensory completion "
        "(Yi et al. 2016), and graph models (Cini et al. 2022; Wu et al. 2019) have improved generic "
        "time-series completion. MixHop convolution aggregates higher-order neighbourhoods (Abu-El-Haija "
        "et al. 2019). Wu et al. (2025) reconstructed Dianchi water-quality fields with biased nonnegative "
        "tensor factorization under random 20–80% holes and 1–4 week gaps, reporting strong RMSE/MAE/"
        "Nash–Sutcliffe scores. That paper answers which reconstructor is most accurate on Dianchi under "
        "random or gap missingness. It does not ask whether the ranking itself depends on the missingness "
        "mechanism, nor whether the same protocol travels to a second aquatic medium. When we score Linear, "
        "LOCF and a MixHop recipe under that bank, Linear wins all eight settings. The same Linear collapses "
        "on Mask-View sensor and station outages (MAE about 0.68). The neighbour’s missingness therefore never "
        "left the interpolable regime. A 2025 Journal of Hydroinformatics review on machine learning for "
        "sensor quality control and imputation likewise concludes that operational trustworthiness remains "
        "difficult to establish (https://doi.org/10.2166/hydro.2025.017).",
    )
    body(
        doc,
        "International water journals increasingly desk-reject single-site machine-learning applications "
        "that use standard architectures without a transferable question (Ahmed et al. 2019). The present "
        "paper is written against that failure mode. We do not claim a new universal sequence model. We "
        "claim that an operational Mask-View bank changes method ranking on (i) an inland automatic network "
        "and (ii) a coastal oxygen forecast cube, and that MCAR-only tables mislead operations. The "
        "forecasting analog is explicit: reconstruction of monthly oxygen maps (Sharp et al. 2023) is not "
        "the same task as a 1–3 month regional forecast under column-limited history (Zheng et al. 2024).",
    )
    body(
        doc,
        "The contributions are: (1) a shared pattern dictionary mapping lake station failures onto "
        "shelf-oxygen masks, including an Argo-column analog; (2) a rank-reversal evaluation in which lake "
        "standardized MAE and ocean RMSE/F1 are compared only through winner identity, including a dual-protocol "
        "table against Wu et al. (2025); (3) evidence that linear interpolation dominates short lake gaps while "
        "collapsing on sensor/station outages, and that a shelf oxygen Transformer is taxed by how much history "
        "remains toward climatology at lead-2, after fair masked persistence is scored.",
    )

    h(doc, "2. Methods")
    h(doc, "2.1. Mask-View protocol", 2)
    body(
        doc,
        "For each window we hide originally observed cells with one of: point (MCAR), block_time "
        "(contiguous downtime), sensor (parameter or depth-channel dropout), station (site or column "
        "outage), mixed, and for the ocean cube argo (profile-like columns) and none (dense-history upper "
        "bound). Training, where applicable, samples complementary views; evaluation uses fixed patterns. "
        "Mask construction differs by medium (Table 1); the scientific question does not.",
    )
    add_table(
        doc,
        ["Pattern", "Lake automatic stations", "Shelf oxygen history"],
        [
            ["point", "MCAR on observed cells", "Static random keep on O2 grid (same voxels all 12 months)"],
            ["block_time", "Consecutive downtime slab", "Time slab × spatial block"],
            ["sensor", "Drop 1–2 variables", "Hide selected O2 depth layers; physics stays"],
            ["station", "Whole station blackout", "Column-limited stations (keep set by column count)"],
            ["mixed", "Point + block (± station)", "Compound of the above"],
            ["argo", "n/a", "Yangtze-section proxy (7 cells); not live Argovis"],
            ["none", "Natural ~19.8% already", "Dense O2 history (upper bound)"],
        ],
    )
    caption(doc, "Table 1. Shared Mask-View pattern bank. White cells in Figure 1 are hidden.")

    h(doc, "2.2. Lake testbed", 2)
    body(
        doc,
        "We use Dianchi-Water: 22 stations, nine variables (water temperature, pH, dissolved oxygen, "
        "conductivity, turbidity, permanganate index, ammonia nitrogen, total phosphorus, total nitrogen), "
        "4-hourly, 2022–2024 (Anonymous 2026). Windows have length 36 (6 days) and stride 18. Temporal split "
        "without leakage: train to 30 June 2023, validation July–December 2023, test year 2024. Features are "
        "standardized on training observed cells. Baselines are Mean, last observation carried forward (LOCF), "
        "linear interpolation per station–variable, SAITS, BRITS, StemGNN reported as GRIN* (official GRIN "
        "is unavailable in PyPOTS 1.5; Du 2026), and MaskView-ST (MixHop graph convolution plus temporal "
        "self-attention). The ablation recipe spatial_plus_l0 uses k = 8 neighbours, 60 km cutoff, "
        "mask-aware aggregation, station-pattern upsampling and consistency weight λ = 0. It is not the "
        "default 105-grid model. Metrics are standardized MAE on artificially hidden, originally observed "
        "cells. Naturally missing cells have no labels; we therefore replay another year’s outage calendar "
        "onto 2024 cells that were observed, with an MCAR control at the same hide count.",
    )

    h(doc, "2.3. Ocean testbed", 2)
    body(
        doc,
        "The East China Sea shelf cube is a WOA-informed development oxygen field with physical drivers "
        "(temperature, salinity, stratification, NOAA OISST, 10 m wind). The task is a 12-month history to "
        "a 1–3 month forecast. Models are persistence, month-of-year climatology, an LSTM anomaly model, a "
        "spatiotemporal Transformer, and a validation-tuned hybrid of climatology and the Transformer "
        "(Vaswani et al. 2017). Masks are applied to oxygen history; physics channels remain visible. "
        "Spatial masks are constant across the 12-month window (except block_time), so a voxel is fully "
        "observed or never. Lake hide rates 10/20/30% are not the same quantity as keep_ratio = 0.25: "
        "station keep is n_stations × Z / n_water (eight columns give 0.089 on this 450-voxel water mask). "
        "Persistence and climatology in the original ablation did not ingest the masked cube; their RMSE "
        "was therefore invariant by construction. Fair counterparts scored here are last-observed persistence "
        "(LOCF) with climatology fill, per-voxel temporal linear interpolation, and horizontal interpolation "
        "of the LOCF field. Climatology remains a valid operational fallback because it does not need recent "
        "oxygen. The learned Transformer does. Uncertainty on the lead-1 Transformer–climatology gap uses a "
        "paired month-block bootstrap (n = 22 test months, 200 resamples) of Δ = RMSEclim − RMSEST. We also "
        "scan point voxel keep at 0.10–0.50 and station column counts 4–24. Sensor is not scanned as a rate: "
        "with Z = 5, 10/20/25% all keep one depth layer. Low-oxygen events use a tenth-percentile threshold. "
        "We do not claim process-level fronts; the cube is not GOBAI-O2 (Sharp et al. 2023).",
    )

    h(doc, "2.4. What is comparable", 2)
    body(
        doc,
        "Lake MAE and ocean RMSE are not on the same scale. We compare (i) which method wins inside each "
        "medium and (ii) whether operational masks flip that ranking relative to the easy-gap or dense-history "
        "case. Forecast verification language follows Wilks (2019) and Jolliffe & Stephenson (2012).",
    )

    h(doc, "3. Results")
    h(doc, "3.1. Pattern bank", 2)
    if maybe_fig(doc, FIGS / "fig_maskview_pattern_bank.png", 6.2):
        caption(
            doc,
            "Figure 1. Mask-View pattern bank. Top: lake windows (time × node×variable). "
            "Bottom: ocean plan-view masks, including Argo-like columns. White = hidden.",
        )

    h(doc, "3.2. Lake ranking is protocol-dependent", 2)
    body(
        doc,
        f"Over {wins.get('n_rows', 105)} lake model–settings, linear interpolation wins "
        f"{wins.get('wins', {}).get('Linear', 9)} individual rate cells, BRITS wins "
        f"{wins.get('wins', {}).get('BRITS', 5)} (all sensor/station), and MaskView-ST wins "
        f"{wins.get('wins', {}).get('MaskView-ST', 1)}. Pattern-averaged standardized MAE is given in Table 2. "
        "Linear is strongest on point, block_time and mixed gaps, reflecting the smoothness of 4-hourly "
        "physicochemistry. A pooled mean MAE would have ranked MaskView-ST first "
        f"({f3(wins['mean_mae_by_model']['MaskView-ST'])} versus Linear "
        f"{f3(wins['mean_mae_by_model']['Linear'])}) and hidden the reversal. On sensor dropout, Linear and "
        "LOCF collapse because no within-series support remains; BRITS recovers cross-variable information. "
        "The spatial_plus_l0 recipe improves sensor MAE from the default MixHop 0.441 to 0.421, which is "
        "0.009 below BRITS 0.430; we treat that as a sensitivity, not a second winner. The station-outage "
        "gap to BRITS remains (~0.105). Ablation with point-only training is excellent on MCAR and collapses "
        "on sensor/station, which is direct evidence that the pattern bank is necessary.",
    )
    lrows = []
    for _, r in lake.iterrows():
        lrows.append(
            [
                r["pattern"],
                f3(r["simple_Linear"]),
                f3(r["learned_MaskViewST"]),
                f3(r["recipe_spatial_plus_l0"]),
                f3(r["BRITS"]),
                r["winner_default_grid"],
            ]
        )
    add_table(
        doc,
        ["Pattern", "Linear", "MaskView-ST", "spatial_plus_l0", "BRITS", "Winner (105-grid)"],
        lrows,
    )
    caption(
        doc,
        "Table 2. Dianchi pattern-averaged standardized MAE (mean of 10/20/30% rates). Lower is better. "
        "Winner is the default 105-grid model; spatial_plus_l0 is an ablation recipe.",
    )

    body(
        doc,
        "The EcoInf 2025 bank does not produce this reversal (Table 3). Linear remains first at random "
        "20/40/60/80% (MAE 0.160, 0.183, 0.209, 0.242) and at 1–4 week contiguous gaps (0.311, 0.285, 0.341, "
        "0.402) against LOCF and spatial_plus_l0. Week-gap Linear sees the full series; the MixHop recipe sees "
        "stitched 6-day windows. Ranking therefore depends on which missingness bank is used, not on swapping "
        "the lake.",
    )
    erows = []
    for proto, setting, label in [
        ("ecoinf_random", "random_20pct", "random 20%"),
        ("ecoinf_random", "random_40pct", "random 40%"),
        ("ecoinf_random", "random_60pct", "random 60%"),
        ("ecoinf_random", "random_80pct", "random 80%"),
        ("ecoinf_week_gap", "1week", "1-week gap"),
        ("ecoinf_week_gap", "2week", "2-week gap"),
        ("ecoinf_week_gap", "3week", "3-week gap"),
        ("ecoinf_week_gap", "4week", "4-week gap"),
    ]:
        erows.append(
            [
                label,
                mae_at(eco, proto, setting, "Linear"),
                mae_at(eco, proto, setting, "LOCF"),
                mae_at(eco, proto, setting, "spatial_plus_l0"),
                "Linear",
            ]
        )
    add_table(doc, ["Setting", "Linear", "LOCF", "spatial_plus_l0", "Winner"], erows)
    caption(
        doc,
        "Table 3. Same Dianchi 2024 test windows under the Wu et al. (2025) missingness bank. "
        "Standardized MAE; Linear wins all eight settings.",
    )

    body(
        doc,
        "Naturally missing cells have no ground truth, so we copy another year’s outage calendar onto 2024 "
        "cells that were observed. Replaying 2022 (hide rate 10.5%) yields Linear 0.383 versus MCAR at the "
        "same count 0.168; replaying 2023 (3.5%) yields 0.328 versus 0.147. Linear still wins both replays; "
        "spatial_plus_l0 does not. Structure therefore taxes Linear relative to MCAR without flipping the "
        "winner. Ranking reversal appears only when the evaluation window itself has no temporal neighbours "
        "(Mask-View sensor/station). The calendar of real outages is interpolable along the year; the "
        "operational stress test is the case with no in-window support.",
    )

    h(doc, "3.3. Ocean: learned skill is taxed by history sparsity", 2)
    body(
        doc,
        "Table 4 reports lead-1 RMSE. Persistence scored on unmasked history is an oracle last-month score "
        "and is not a sparse baseline. persist_locf uses only kept oxygen. persist_locf never beats "
        "climatology (5.30 µmol kg−1). Temporal linear interpolation equals persist_locf under static spatial "
        "masks because a last-month forecast only needs the endpoint. Horizontal interpolation of sparse "
        "columns is worse than persistence (RMSE 12–15). The competitive simple bar is therefore climatology. "
        "The Transformer is best under dense history (3.88; low-oxygen F1 0.74) and remains first at lead-1 "
        "under every operational mask, but RMSE rises 10% (block_time) to 35% (station/argo) as voxel keep "
        "falls. At lead-2 the best model is climatology or a hybrid except for block_time. Low-oxygen F1 falls "
        "from 0.74 (dense) to 0.67 (argo). Argo and station differ by 0.009 µmol kg−1 and share a column-limited "
        "geometry; argo here is a Yangtze-section proxy, not live Argovis.",
    )
    orows = []
    O = ocean.set_index("pattern")
    for pat in ["none", "block_time", "point", "sensor", "station", "argo"]:
        r = O.loc[pat]
        deg = "—" if pat == "none" else f"+{100 * float(r['degradation_vs_dense']):.0f}%"
        orows.append(
            [
                pat,
                f3(r["keep_frac"]),
                f3(r["persist_locf"]),
                f3(r["clim_RMSE"]),
                f3(r["ST_RMSE"]),
                deg,
                f3(r["st_margin_vs_clim"]),
            ]
        )
    add_table(
        doc,
        ["Pattern", "Keep", "persist_locf", "Climatology", "ST", "ST vs dense", "clim − ST"],
        orows,
    )
    caption(
        doc,
        "Table 4. East China Sea lead-1 RMSE (µmol kg−1). persist_locf ingests the same mask; "
        "climatology does not need recent oxygen. Locked ablation ST RMSE.",
    )

    bnone = boot[boot.pattern == "none"].iloc[0]
    bpt = boot[boot.pattern == "point"].iloc[0]
    bstn = boot[boot.pattern == "station"].iloc[0]
    bargo = boot[boot.pattern == "argo"].iloc[0]
    body(
        doc,
        "The remaining station/section edge looks small enough to be noise if one only inspects overlapping "
        "marginal RMSE intervals. The paired difference is the right test. Month-resampled Δ = clim − ST stays "
        f"positive in all 200 replicates: dense history {f3(bnone.point_delta_clim_minus_st)} "
        f"[{f3(bnone.delta_clim_minus_st_p05)}, {f3(bnone.delta_clim_minus_st_p95)}]; "
        f"point {f3(bpt.point_delta_clim_minus_st)} "
        f"[{f3(bpt.delta_clim_minus_st_p05)}, {f3(bpt.delta_clim_minus_st_p95)}]; "
        f"station {f3(bstn.point_delta_clim_minus_st)} "
        f"[{f3(bstn.delta_clim_minus_st_p05)}, {f3(bstn.delta_clim_minus_st_p95)}]; "
        f"section {f3(bargo.point_delta_clim_minus_st)} "
        f"[{f3(bargo.delta_clim_minus_st_p05)}, {f3(bargo.delta_clim_minus_st_p95)}]. "
        "Retrain RMSE tracks the locked ablation to within 0.02 µmol kg−1.",
    )

    body(
        doc,
        "That default table confounds rate with geometry: station keeps 0.089 of water voxels, point keeps "
        "0.25. A keep scan (Figure 3, Table 5) separates the two. On point (the lake 10/20/30% analog) the "
        "Transformer wins every rate. Matching voxel keep, columns are only modestly harsher. Four columns "
        "(keep 0.044) shrink Δ to 0.013 but the paired interval still excludes zero. Operational sparsity "
        "therefore does not invert lead-1 ranking. It taxes the learned advantage mainly by how much oxygen "
        "history remains, secondarily by packing that history into columns, and it hands longer leads back "
        "to climatology. The residual 0.06 µmol kg−1 at the default eight-column mask is statistically signed "
        "on this window, not operationally large.",
    )
    if maybe_fig(doc, FIGS / "fig_keep_ratio_tax.png", 5.6):
        caption(
            doc,
            "Figure 2. Lead-1 Δ = climatology RMSE − Transformer RMSE versus effective voxel keep. "
            "Error bars are paired month-block 5–95% intervals (n = 22, 200 resamples). "
            "Point is the lake 10/20/30% analog; station varies column count.",
        )
    krows = [
        ["point 10%", "point"] + keep_row(keep, pattern="point", keep_ratio=0.1),
        ["point 20%", "point"] + keep_row(keep, pattern="point", keep_ratio=0.2),
        ["point 30%", "point"] + keep_row(keep, pattern="point", keep_ratio=0.3),
        ["4 columns", "station"] + keep_row(keep, pattern="station", n_stations=4),
        ["8 columns (default)", "station"] + keep_row(keep, pattern="station", n_stations=8),
        ["16 columns", "station"] + keep_row(keep, pattern="station", n_stations=16),
        ["24 columns", "station"] + keep_row(keep, pattern="station", n_stations=24),
    ]
    add_table(doc, ["Setting", "Geometry", "Keep", "ST", "Δ (clim−ST)", "Δ 5–95%"], krows)
    caption(
        doc,
        "Table 5. Keep scan, 8-epoch ablation recipe. Δ is climatology RMSE minus Transformer RMSE "
        "(µmol kg−1). ST remains first in every row.",
    )

    h(doc, "3.4. Cross-media message", 2)
    if maybe_fig(doc, FIGS / "fig_rank_reversal_two_media.png", 6.3):
        caption(
            doc,
            "Figure 3. Rankings under the shared bank. Left: Dianchi MAE for Linear, MaskView-ST and "
            "spatial_plus_l0. Right: ECS lead-1 RMSE for persist_locf, climatology and the Transformer.",
        )
    body(
        doc,
        "Easy lake gaps reward Linear; hard lake gaps reward cross-variable and spatial models. Real Dianchi "
        "outage calendars still rank Linear first but roughly double its MAE versus MCAR at the same hide "
        "rate. Dense ocean history rewards a Transformer; column-limited history does not let persist/linear "
        "steal the win. The published station versus point contrast (+35% versus +29% versus dense) is mostly "
        "that default station keeps about 9% of voxels versus point 25%. Lake-like point 10/20/30% never flips "
        "lead-1 ranking. Lead-2 reverts to climatology. A benchmark that only injects point holes, or that "
        "scores persistence on unmasked history, would have missed both failure modes. Table 6 records winner "
        "identity on shared pattern names.",
    )
    rrows = []
    for _, r in rr.iterrows():
        rrows.append(
            [
                r["pattern"],
                r["lake_winner"],
                r["ocean_lead1_best"],
                r["ocean_lead2_best"],
                f3(r["ocean_lead1_degradation"]),
            ]
        )
    add_table(
        doc,
        ["Pattern", "Lake winner", "Ocean lead-1 best", "Ocean lead-2 best", "Ocean ST degradation"],
        rrows,
    )
    caption(doc, "Table 6. Shared-pattern winner identity. Degradation is relative to dense-history ST RMSE.")

    h(doc, "4. Discussion")
    body(
        doc,
        "Wu et al. (2025) remain the reconstruction-accuracy neighbour; we differ in question (ranking under "
        "operational masks; two media). MixHop capacity does not close the lake station-outage gap to BRITS, "
        "which mixes flattened cross-station channels more aggressively than a local graph "
        "(Kipf & Welling 2017; Veličković et al. 2018). That gap is informative rather than embarrassing.",
    )
    body(
        doc,
        "Ocean oxygen here is a WOA-informed development cube. Ranking under masks still holds as a protocol "
        "demonstration; hypoxia early-warning products for the East China Sea would need time-varying oxygen "
        "targets and, for short-range bays, hydrodynamic predictors (Zheng et al. 2024). Unmasked persistence/"
        "climatology invariance must not be reported as a scientific finding; it follows from baselines that "
        "ignore the mask. After fair scoring, climatology remains the simple method to beat. The Transformer’s "
        "remaining lead-1 edge under the default eight-column mask is small because that mask keeps about 9% "
        "of voxels, not because columns uniquely destroy skill. On n = 22 test months the paired bootstrap of "
        "Δ excludes zero down to four columns. That supports still first as a ranking statement on this window; "
        "it does not make 0.06 µmol kg−1 an operationally large gain, and month exchangeability is an assumption "
        "of the block bootstrap. Lake 10/20/30% should be compared with ocean point keep, not with keep_ratio = "
        "0.25 on station.",
    )
    body(
        doc,
        "We recommend that hydroinformatics studies (i) publish the pattern bank, (ii) keep linear interpolation "
        "and climatology as first-class baselines, and (iii) report winner identity by pattern rather than a "
        "single mean score.",
    )

    h(doc, "5. Conclusions")
    body(
        doc,
        "Mask-View is an evaluation and training protocol for operational aquatic missingness. On Dianchi it "
        "shows when not to prefer a graph imputer over linear interpolation; on the East China Sea shelf it "
        "shows when a Transformer’s dense-history advantage is taxed back toward climatology. Missingness "
        "protocol is part of the scientific claim.",
    )

    h(doc, "Data availability statement")
    body(
        doc,
        "Frozen tables and figures are in this repository (https://github.com/Az0998/maskview-aquatic-protocol). "
        "Lake experiment engine: https://github.com/Az0998/dianchi-maskview-imputation. Ocean experiment engine: "
        "https://github.com/Az0998/ocean-do-forecast. Dianchi-Water data: "
        "https://huggingface.co/datasets/anonymous-dianchi-2026/dianchi-water (CC BY 4.0; accessed 8 August 2026).",
    )
    h(doc, "Author contributions")
    body(
        doc,
        "S.Z. conceived the study, designed the Mask-View protocol, performed the analyses, and wrote the manuscript.",
    )
    h(doc, "Acknowledgements")
    body(
        doc,
        "The author thanks the compilers of the open Dianchi monitoring dataset and the China National "
        "Environmental Monitoring Center data-release system.",
    )
    h(doc, "Conflict of interest")
    body(doc, "The author declares no conflict of interest.")
    h(doc, "Funding")
    body(doc, "This research received no external funding.")

    h(doc, "References")
    refs = [
        "Abu-El-Haija S., Perozzi B., Kapoor A., Alipourfard N., Lerman K., Harutyunyan H., Ver Steeg G. & Galstyan A. 2019 MixHop: higher-order graph convolutional architectures via sparsified neighborhood mixing. In: Proceedings of the 36th International Conference on Machine Learning, pp. 21–29.",
        "Ahmed A. N., Othman F. B., Afan H. A., Ibrahim R. K., Fai C. M., Hossain M. S., Ehteram M. & Elshafie A. 2019 Machine learning methods for better water quality prediction. Journal of Hydrology 578, 124084.",
        "Anonymous 2026 dianchi-water: high-frequency multi-station surface water quality dataset for the Dianchi Lake basin (2022–2024). Hugging Face Datasets. https://huggingface.co/datasets/anonymous-dianchi-2026/dianchi-water (accessed 8 August 2026).",
        "Breitburg D., Levin L. A., Oschlies A., Grégoire M., Chavez F. P., Conley D. J., Garçon V., Gilbert D., Gutiérrez D., Isensee K., Jacinto G. S., Limburg K. E., Montes I., Naqvi S. W. A., Pitcher G. C., Rabalais N. N., Roman M. R., Rose K. A., Seibel B. A., Telszewski M., Yasuhara M. & Zhang J. 2018 Declining oxygen in the global ocean and coastal waters. Science 359 (6371), eaam7240.",
        "Cao W., Wang D., Li J., Zhou H., Li L. & Li Y. 2018 BRITS: bidirectional recurrent imputation for time series. In: Advances in Neural Information Processing Systems 31.",
        "Che Z., Purushotham S., Cho K., Sontag D. & Liu Y. 2018 Recurrent neural networks for multivariate time series with missing values. Scientific Reports 8, 6085.",
        "China National Environmental Monitoring Center (CNEMC) 2026 National surface water quality automatic monitoring data release system. https://szzdjc.cnemc.cn (accessed 8 August 2026).",
        "Cini A., Marisca I. & Alippi C. 2022 Filling the GApS: multivariate time series imputation by graph neural networks. In: International Conference on Learning Representations.",
        "Diaz R. J. & Rosenberg R. 2008 Spreading dead zones and consequences for marine ecosystems. Science 321 (5891), 926–929.",
        "Du W. 2026 PyPOTS: a Python toolbox for data mining on partially-observed time series. https://pypots.com (accessed 8 August 2026).",
        "Du W., Côté D. & Liu Y. 2023 SAITS: self-attention-based imputation for time series. Expert Systems with Applications 219, 119619.",
        "Gao W., Howarth R. W., Swaney D. P., Hong B. & Guo H. C. 2015 Enhanced N input to Lake Dianchi Basin from 1980 to 2010: drivers and implications. Environmental Pollution 198, 234–242.",
        "Huang C., Wang X., Yang H., Li Y., Wang Y., Chen X. & Xu L. 2014 Satellite data regarding the eutrophication response to human activities in the plateau lake Dianchi in China from 1974 to 2009. Science of the Total Environment 485–486, 1–11.",
        "IWA Publishing 2026 Instructions for authors. https://iwaponline.com/pages/Instructions_for_authors (accessed 26 August 2026).",
        "Jolliffe I. T. & Stephenson D. B. 2012 Forecast Verification: A Practitioner’s Guide in Atmospheric Science, 2nd edn. Wiley, Chichester, UK.",
        "Journal of Hydroinformatics 2025 A review on how machine learning can be beneficial for sensor data quality control and imputation in water resources management. Journal of Hydroinformatics. https://doi.org/10.2166/hydro.2025.017.",
        "Kipf T. N. & Welling M. 2017 Semi-supervised classification with graph convolutional networks. In: International Conference on Learning Representations.",
        "Little R. J. A. & Rubin D. B. 2019 Statistical Analysis with Missing Data, 3rd edn. Wiley, Hoboken, NJ, USA.",
        "Luo Y., Cai X., Zhang Y., Xu J. & Yuan X. 2018 Multivariate time series imputation with generative adversarial networks. In: Advances in Neural Information Processing Systems 31.",
        "Ministry of Ecology and Environment of the People’s Republic of China (MEE) 2002 Environmental Quality Standards for Surface Water (GB 3838-2002). Beijing, China.",
        "Paerl H. W. & Huisman J. 2008 Blooms like it hot. Science 320 (5872), 57–58.",
        "Sharp J. D., Fassbender A. J., Carter B. R., Johnson G. C., Schultz C. & Dunne J. P. 2023 GOBAI-O2: temporally and spatially resolved fields of ocean interior dissolved oxygen over nearly two decades. Earth System Science Data 15, 4481–4518.",
        "Tashiro Y., Song J., Song Y. & Ermon S. 2021 CSDI: conditional score-based diffusion models for probabilistic time series imputation. In: Advances in Neural Information Processing Systems 34.",
        "van Buuren S. 2018 Flexible Imputation of Missing Data, 2nd edn. CRC Press, Boca Raton, FL, USA.",
        "Vaswani A., Shazeer N., Parmar N., Uszkoreit J., Jones L., Gomez A. N., Kaiser Ł. & Polosukhin I. 2017 Attention is all you need. In: Advances in Neural Information Processing Systems 30.",
        "Veličković P., Cucurull G., Casanova A., Romero A., Liò P. & Bengio Y. 2018 Graph attention networks. In: International Conference on Learning Representations.",
        "Wilks D. S. 2019 Statistical Methods in the Atmospheric Sciences, 4th edn. Elsevier, Amsterdam, The Netherlands.",
        "Wu X., Shan K., Wang L., Wang J. & Shang M. 2025 Spatiotemporal water quality data reconstruction: a tensor factorization framework. Ecological Informatics 90, 103283. https://doi.org/10.1016/j.ecoinf.2025.103283.",
        "Wu Z., Pan S., Long G., Jiang J. & Zhang C. 2019 Graph WaveNet for deep spatial-temporal graph modeling. In: Proceedings of the 28th International Joint Conference on Artificial Intelligence, pp. 1907–1913.",
        "Yi X., Zheng Y., Zhang J. & Li T. 2016 ST-MVL: filling missing values in geo-sensory time series data. In: Proceedings of the 25th International Joint Conference on Artificial Intelligence, pp. 2704–2710.",
        "Zhang Y., Ma R., Zhang M., Duan H., Loiselle S. & Xu J. 2015 Fourteen-year record (2000–2013) of the spatial and temporal dynamics of floating algae blooms in Lake Dianchi, China. Remote Sensing 7, 11203–11222.",
        "Zheng G., Li M., Harding L. W. Jr. & DiGiacomo P. M. 2024 Hypoxia forecasting for Chesapeake Bay using artificial intelligence. Artificial Intelligence for the Earth Systems 3 (3), e230054.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        font(p.add_run(r), size=10)
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.35)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15

    PAPER.mkdir(parents=True, exist_ok=True)
    out = PAPER / "JHI_MaskView_protocol_manuscript.docx"
    try:
        doc.save(out)
    except PermissionError:
        out = PAPER / "JHI_MaskView_protocol_manuscript_v2.docx"
        doc.save(out)
        print("NOTE: dest locked; wrote", out)
    print("Wrote", out, "n_refs=", len(refs), "n_highlights=", len(highlights))

    cover = Document()
    cover.styles["Normal"].font.name = "Times New Roman"
    cover.styles["Normal"].font.size = Pt(11)
    for para in (PAPER / "cover_letter_jhi.txt").read_text(encoding="utf-8").splitlines():
        p = cover.add_paragraph()
        font(p.add_run(para if para else " "), size=11)
        p.paragraph_format.space_after = Pt(0 if para else 8)
        p.paragraph_format.line_spacing = 1.15
    cover_out = PAPER / "cover_letter_jhi.docx"
    cover.save(cover_out)
    print("Wrote", cover_out)


if __name__ == "__main__":
    main()

